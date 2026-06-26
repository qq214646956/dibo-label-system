"""Convert 用户操作手册.md to Word document."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

doc = Document()

# Override the default font at the style XML level so ALL runs inherit it
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
# Also set the East-Asian font
rpr = style.element.get_or_add_rPr()
rFonts = rpr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = rpr.makeelement(qn('w:rFonts'), {})
    rpr.insert(0, rFonts)
rFonts.set(qn('w:ascii'), '微软雅黑')
rFonts.set(qn('w:hAnsi'), '微软雅黑')
rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Also override heading styles
for level in [1, 2, 3]:
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    rpr = hs.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), '微软雅黑')
    rFonts.set(qn('w:hAnsi'), '微软雅黑')
    rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Override list styles
for name in ['List Bullet', 'List Number']:
    try:
        ls = doc.styles[name]
        ls.font.name = '微软雅黑'
        rpr = ls.element.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rpr.makeelement(qn('w:rFonts'), {})
            rpr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), '微软雅黑')
        rFonts.set(qn('w:hAnsi'), '微软雅黑')
        rFonts.set(qn('w:eastAsia'), '微软雅黑')
    except KeyError:
        pass

def set_run_font(run, name='微软雅黑', size=None, color=None, bold=None, italic=None):
    """Explicitly set font on a run to avoid Calibri bleed."""
    run.font.name = name
    rpr = run._r.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.font.italic = italic

# Read markdown
with open(r'D:\我的文件\地博外部项目\10标签打印系统\标签打印系统\用户操作手册.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

in_table = False
table_rows = []

def flush_table():
    global in_table, table_rows
    if not table_rows:
        return
    ncols = max(len(r) for r in table_rows)
    tbl = doc.add_table(rows=len(table_rows), cols=ncols, style='Light Grid Accent 1')
    for i, row in enumerate(table_rows):
        for j, cell_text in enumerate(row):
            cell = tbl.cell(i, j)
            cell.text = cell_text
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=Pt(9))
    doc.add_paragraph()
    table_rows = []
    in_table = False

for line in lines:
    line = line.rstrip()

    if not line:
        if in_table:
            flush_table()
        continue

    # Image placeholders
    if line.startswith('> 【插图'):
        p = doc.add_paragraph()
        run = p.add_run(line[2:])
        set_run_font(run, size=Pt(9), color=RGBColor(128, 128, 128), italic=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        continue

    # H1
    m = re.match(r'^# (.+)', line)
    if m:
        flush_table()
        h = doc.add_heading(m.group(1), level=1)
        for run in h.runs:
            set_run_font(run)
        continue

    # H2
    m = re.match(r'^## (.+)', line)
    if m:
        flush_table()
        h = doc.add_heading(m.group(1), level=2)
        for run in h.runs:
            set_run_font(run)
        continue

    # H3
    m = re.match(r'^### (.+)', line)
    if m:
        flush_table()
        h = doc.add_heading(m.group(1), level=3)
        for run in h.runs:
            set_run_font(run)
        continue

    # Table separator
    if re.match(r'^\|[- |]+\|$', line):
        continue

    # Table row
    if line.startswith('|') and line.endswith('|'):
        cells = [c.strip() for c in line[1:-1].split('|')]
        table_rows.append(cells)
        in_table = True
        continue

    # Bullet list
    m = re.match(r'^- (.+)', line)
    if m:
        flush_table()
        p = doc.add_paragraph(m.group(1), style='List Bullet')
        for run in p.runs:
            set_run_font(run)
        continue

    # Numbered list
    m = re.match(r'^\d+\. (.+)', line)
    if m:
        flush_table()
        p = doc.add_paragraph(m.group(1), style='List Number')
        for run in p.runs:
            set_run_font(run)
        continue

    # Code block
    if line.startswith('`'):
        flush_table()
        code = line.strip('`')
        p = doc.add_paragraph()
        run = p.add_run(code)
        set_run_font(run, name='Consolas', size=Pt(9))
        p.paragraph_format.left_indent = Cm(1)
        continue

    # Bold text
    flush_table()
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*.*?\*\*)', line)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = p.add_run(part)
            set_run_font(run)

flush_table()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

out_path = r'D:\我的文件\地博外部项目\10标签打印系统\标签打印系统\用户操作手册.docx'
doc.save(out_path)
print(f'Done: {out_path}')
